"""Publication figures (JPEG) and JHSS-formatted .docx."""

from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import statsmodels.formula.api as smf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CLEAN = ROOT / "data" / "clean"
FIG = ROOT / "paper" / "figures"
DOCX = ROOT / "paper" / "Lagged_pack_depth_DCSAA_boys_XC.docx"

NAVY = "#1B365D"
TEAL = "#2A6F7F"
GOLD = "#C4A35A"
GRAY = "#4A4A4A"


def _style():
    plt.rcParams.update(
        {
            "font.family": "serif",
            "font.size": 10,
            "axes.titlesize": 11,
            "axes.labelsize": 10,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "axes.edgecolor": GRAY,
            "axes.labelcolor": "#222222",
            "xtick.color": "#222222",
            "ytick.color": "#222222",
            "axes.spines.top": False,
            "axes.spines.right": False,
        }
    )


def load_fit() -> tuple[pd.DataFrame, pd.DataFrame]:
    panel = pd.read_csv(CLEAN / "team_seasons.csv")
    fit = panel.loc[panel["fit_set"] & panel["n_finishers"].ge(5)].copy()
    hold = panel.loc[panel["holdout"] & panel["n_finishers"].ge(5)].copy()
    return fit, hold


def make_figures() -> dict[str, Path]:
    _style()
    FIG.mkdir(parents=True, exist_ok=True)
    fit, hold = load_fit()
    audit = pd.read_csv(CLEAN / "season_audit.csv")
    paths = {}
    colors = {"private": NAVY, "public": TEAL, "charter": GOLD}

    fig, ax = plt.subplots(figsize=(6.4, 3.6))
    x = np.arange(len(audit))
    w = 0.27
    ax.bar(x - w, audit["individual_finishers_parsed"], width=w, color=GOLD, label="Individual finishers parsed")
    ax.bar(x, audit["scoring_teams_ge5"], width=w, color=TEAL, label="Scoring teams (≥5)")
    ax.bar(x + w, audit["estimation_rows"] + audit["holdout_rows"], width=w, color=NAVY, label="Lagged analytic rows")
    ax.set_xticks(x)
    ax.set_xticklabels(audit["season"].astype(int).astype(str), rotation=0)
    ax.set_xlabel("Season")
    ax.set_ylabel("Count")
    ax.set_title("Season-by-season inclusion")
    ax.legend(frameon=False, fontsize=8)
    fig.tight_layout()
    p = FIG / "figure1_inclusion.jpg"
    fig.savefig(p, dpi=300, bbox_inches="tight")
    plt.close(fig)
    paths["f1"] = p

    d = fit.dropna(subset=["lag_depth_z", "state_place", "lag_pack_gap", "lag_star_z"])
    fig, ax = plt.subplots(figsize=(6.2, 4.3))
    for sector, g in d.groupby("sector"):
        ax.scatter(
            g["lag_pack_gap"],
            g["state_place"],
            c=colors.get(sector, GRAY),
            label=sector.title(),
            s=42,
            alpha=0.85,
            edgecolors="white",
            linewidths=0.4,
        )
    band = pd.DataFrame({"x": d["lag_pack_gap"], "y": d["state_place"]})
    m = smf.ols("y ~ x", band).fit()
    xs = np.linspace(band["x"].min(), band["x"].max(), 60)
    pr = m.get_prediction(pd.DataFrame({"x": xs})).summary_frame()
    ax.plot(xs, pr["mean"], color=GRAY, lw=1.3)
    ax.fill_between(xs, pr["mean_ci_lower"], pr["mean_ci_upper"], color=GRAY, alpha=0.18)
    ax.set_xlabel("Prior-year pack gap (fifth-runner z − first-runner z)")
    ax.set_ylabel("Next-year championship team place (lower is better)")
    ax.invert_yaxis()
    ax.legend(frameon=False)
    fig.tight_layout()
    p = FIG / "figure2_pack_gap.jpg"
    fig.savefig(p, dpi=300, bbox_inches="tight")
    plt.close(fig)
    paths["f2"] = p

    xlim = (
        min(d["lag_star_z"].min(), d["lag_depth_z"].min()) - 0.15,
        max(d["lag_star_z"].max(), d["lag_depth_z"].max()) + 0.15,
    )
    ylim = (d["state_place"].max() + 0.6, 0.4)
    fig, axes = plt.subplots(1, 2, figsize=(7.2, 3.8), sharey=True)
    panels = [
        (axes[0], "lag_star_z", "Prior-year first-runner z", NAVY),
        (axes[1], "lag_depth_z", "Prior-year fifth-runner z", GOLD),
    ]
    for ax, col, title, c in panels:
        ax.scatter(d[col], d["state_place"], c=c, s=36, alpha=0.85, edgecolors="white", linewidths=0.4)
        band = pd.DataFrame({"x": d[col], "y": d["state_place"]})
        m = smf.ols("y ~ x", band).fit()
        xs = np.linspace(band["x"].min(), band["x"].max(), 60)
        pr = m.get_prediction(pd.DataFrame({"x": xs})).summary_frame()
        ax.plot(xs, pr["mean"], color=GRAY, lw=1.2)
        ax.fill_between(xs, pr["mean_ci_lower"], pr["mean_ci_upper"], color=GRAY, alpha=0.16)
        ax.set_xlim(xlim)
        ax.set_ylim(ylim)
        ax.set_xlabel(title)
        ax.set_title(title.replace(" z", ""))
    axes[0].set_ylabel("Next-year team place")
    fig.tight_layout()
    p = FIG / "figure3_star_depth.jpg"
    fig.savefig(p, dpi=300, bbox_inches="tight")
    plt.close(fig)
    paths["f3"] = p

    m = smf.ols(
        "state_place ~ log_enrollment + C(sector) + lag_depth_z",
        data=fit,
    ).fit(cov_type="cluster", cov_kwds={"groups": fit["school"], "use_correction": True})
    hold = hold.dropna(subset=["lag_depth_z", "log_enrollment", "sector"]).copy()
    hold["pred"] = m.predict(hold)
    fig, ax = plt.subplots(figsize=(6.4, 5.0))
    ax.scatter(hold["state_place"], hold["pred"], c=TEAL, s=46, zorder=3)
    lims = [0.3, max(hold["state_place"].max(), hold["pred"].max()) + 1.2]
    ax.plot(lims, lims, color=GRAY, ls="--", lw=1)
    for _, r in hold.iterrows():
        label = str(r["school"]).replace(" College", "").replace(" School", "")
        if len(label) > 16:
            label = label[:15] + "."
        ax.annotate(label, (r["state_place"], r["pred"]), fontsize=6.5, xytext=(3, 3), textcoords="offset points")
    ax.set_xlabel("Observed 2025 team place")
    ax.set_ylabel("Predicted 2025 team place (depth model)")
    ax.set_xlim(lims)
    ax.set_ylim(lims)
    ax.set_aspect("equal", adjustable="box")
    fig.tight_layout()
    p = FIG / "figure4_holdout.jpg"
    fig.savefig(p, dpi=300, bbox_inches="tight")
    plt.close(fig)
    paths["f4"] = p
    return paths


def _set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_blank(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Inches(0)
    return p


def add_para(doc: Document, text: str, *, size=12, bold=False, italic=False, center=False, space_after_blank=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.first_line_indent = Inches(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    if space_after_blank:
        add_blank(doc)
    return p


def add_heading_line(doc: Document, text: str):
    add_para(doc, text, bold=True, space_after_blank=True)


def add_mixed(doc: Document, parts: list[tuple[str, dict]], space_after_blank=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    pf.first_line_indent = Inches(0)
    for text, kw in parts:
        run = p.add_run(text)
        _set_run_font(run, **kw)
    if space_after_blank:
        add_blank(doc)
    return p


def shade_header(cell, hex_color="1B365D"):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str):
    add_para(doc, caption.split(". ", 1)[0] + ".", bold=True, size=12, space_after_blank=False)
    rest = caption.split(". ", 1)[1] if ". " in caption else ""
    if rest:
        add_para(doc, rest, size=12, space_after_blank=True)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        _set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        shade_header(hdr[i])
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(val)
            _set_run_font(run, size=10)
    add_blank(doc)


def add_figure(doc: Document, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.0))
    cap = doc.add_paragraph()
    cap.paragraph_format.first_line_indent = Inches(0)
    cap.paragraph_format.space_after = Pt(0)
    cap.paragraph_format.line_spacing = 1.15
    r = cap.add_run(caption)
    _set_run_font(r, size=10)
    add_blank(doc)


def build_docx(paths: dict[str, Path]) -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.first_line_indent = Inches(0)

    add_para(
        doc,
        "Star Power or Pack Depth? Predicting DCSAA Boys’ Cross-Country Team Place from Prior-Year Results",
        size=14,
        bold=True,
        center=True,
    )
    add_para(doc, "Original article", italic=True, center=True, size=12)

    add_heading_line(doc, "Abstract")
    add_para(
        doc,
        "Background. High-school cross-country team scores are the sum of the first five finishing places, so a school’s fifth scorer can decide a championship as surely as its fastest runner. Coaches often contrast “star power” with “pack depth,” but those labels have rarely been tested as lagged forecasts on a fixed championship course. Objective. To determine whether prior-year fifth-runner performance predicts next-year District of Columbia State Athletic Association (DCSAA) boys’ varsity team place more usefully than prior-year first-runner performance after enrollment and school sector are accounted for. Methods. Public Kenilworth Park championship files from 2019 and 2021–2025 were parsed. Within each meet, first- and fifth-runner times were converted to field-wide z-scores. Nested ordinary least-squares models of team place and team score used school-clustered standard errors, a wild cluster bootstrap, leave-one-school-out prediction, and a frozen 2025 temporal holdout. Pack gap (fifth-runner z minus first-runner z) entered the nested set. The same lagged design was repeated on girls’ files as a replication. Results. The boys’ estimation sample comprised 40 team-seasons from 15 schools (holdout n = 12). Controls explained 46.5% of the variance in team place. Adding first-runner z raised R² to 0.62; adding fifth-runner z raised R² to 0.69 and produced the lowest Akaike information criterion. In the joint model, fifth-runner z remained associated with place while first-runner z did not. Leave-one-school-out RMSE was 2.80 places. On 2025 holdout teams, the depth model had Spearman ρ = 0.87 and RMSE = 3.51; controls had RMSE = 3.08; pack gap had RMSE = 2.53. Girls’ models (n = 18) did not rank depth above first-runner z. Conclusions. In this small DCSAA boys’ panel, prior-year fifth-runner form was the stronger in-sample predictor of next-year team place. Holdout RMSE and the girls’ replication do not support treating pack depth as a general forecasting rule. The associations are not causal.",
    )
    add_mixed(
        doc,
        [
            ("Keywords. ", {"size": 12, "bold": True, "italic": False}),
            (
                "Boys’ cross-country, DCSAA, Kenilworth Park, team place, pack depth, first-runner performance, within-meet z-score, lagged predictors, cluster-robust standard errors, holdout validation, school sector, enrollment, Akaike information criterion, wild cluster bootstrap, high school sports",
                {"size": 12, "italic": True},
            ),
        ],
    )
    add_para(doc, "[Author Name]", center=True)
    add_para(doc, "[High School], [Street Address], [City], [State ZIP], USA", center=True, size=11)
    add_para(doc, "[email]", center=True, italic=True, size=11)

    add_heading_line(doc, "1. Introduction")
    add_para(
        doc,
        "Under National Federation of State High School Associations rules, a cross-country team’s score is the sum of the finishing places of its first five runners (1). Displacement of a fifth scorer by a few places can therefore reverse a dual meet or a championship. That arithmetic motivates a coaching contrast that is widely discussed and rarely measured. Star power treats the fastest athlete as the best summary of a program. Pack depth treats the compactness of the scoring five, or the quality of the fifth scorer, as the better summary. Because each of the first five places counts equally, the second claim is at least as consistent with the scoring rule as the first.",
    )
    add_para(
        doc,
        "The contrast is easy to test poorly. Same-meet first- and fifth-runner times are mechanical ingredients of that meet’s team score. Using them to “predict” the same day’s place is nearly tautological. A coach who wants a forecast, however, already has last year’s championship results before this year’s race. The scientifically usable question is therefore lagged: after school size and sector are accounted for, does last year’s fifth-runner form predict this year’s team place more closely than last year’s first-runner form?",
    )
    add_para(
        doc,
        "Sports rating research has long used least squares to adjust records for opponent strength (2, 3). Related work has used schedule metrics in college basketball (4) and Monte Carlo simulation for high-school football seeding (5). Public–private competitive balance is a documented feature of interscholastic sport (6). Those studies address ranking systems, playoff design, or institutional equity. They do not compare lagged first-runner versus fifth-runner performance as forecasts of a high-school cross-country championship on a single course.",
    )
    add_para(
        doc,
        "The District of Columbia State Athletic Association (DCSAA) boys’ varsity race is a useful test bed for that comparison (7). All classifications contest one championship rather than separate state meets. The varsity course at Kenilworth Park has been 5 km, with a usual seven-runner cap and a five-finisher requirement to score (7). Because the venue is fixed, within-meet z-scores absorb year-to-year differences in weather and field quality without a multi-course conversion. Sector and enrollment remain the relevant background variables in the absence of classification splits.",
    )
    add_para(
        doc,
        "The primary hypothesis was that, after log enrollment and school sector, prior-year fifth-runner within-meet z-score would predict next-year DCSAA boys’ team place more usefully than prior-year first-runner z-score. Team score was treated as a co-primary continuous outcome because it is the quantity the rules actually sum. Pack gap, defined as fifth-runner z minus first-runner z, was included as an operationalization of spread beyond the front of the team. Washington Catholic Athletic Conference (WCAC) membership was an exploratory proxy for regular-season schedule strength, not a co-equal hypothesis. Invitational strength-of-schedule files were not available. A girls’ panel using the same lagged Kenilworth design served as a replication, not as a pooled sample. The analysis does not establish that changing fifth-runner performance would cause a change in team place.",
    )

    add_heading_line(doc, "2. Materials and Methods")
    add_mixed(doc, [("2.1 Design and outcomes", {"size": 12, "bold": True, "italic": True})], space_after_blank=False)
    add_para(
        doc,
        "The unit of analysis was the school–season. Team place at the DCSAA championship was a primary outcome because that is the ranking the meet publishes and the quantity a coach would forecast. Place is an ordinal, bounded ranking; ordinary least squares is used as a linear approximation for a short ranking, with school-clustered standard errors and a finite-sample correction (8). Team score, the sum of the first five places, was a co-primary continuous outcome that matches the scoring rule more closely. Predictors were taken only from the previous available Kenilworth championship. The 2020 season had no championship, so 2021 was linked to 2019. After the 2023 individual file was recovered, 2024 was linked to 2023. The 2025 season was reserved as a single temporal holdout and was not used to choose specifications.",
    )

    add_mixed(doc, [("2.2 Data sources and authority", {"size": 12, "bold": True, "italic": True})], space_after_blank=False)
    add_para(
        doc,
        "Meet identifiers were taken from the public MileSplit championship index and, for 2025, from M&D Timing (Table 1) (7). Local copies of result pages are stored with the analysis code. When a posted team-score table and a reconstructed table both existed, the posted table was treated as authoritative. Individual times were used to construct first- and fifth-runner z-scores and to reconstruct scores when no team table was present. Reconstructed ranks ordered schools by reconstructed score and then by school name.",
    )
    add_para(
        doc,
        "School strings were mapped through a directory of aliases, sector (private, public, charter), WCAC membership, and a static grades 9–12 enrollment estimate. Matching used a normalized alias table. Unmatched names were dropped. Enrollment is not a year-by-year official count; it is treated as known independently of the outcome year. WCAC membership was coded for Gonzaga College, St. John’s College, and Archbishop Carroll.",
    )

    add_table(
        doc,
        ["Season", "File", "Finishers parsed", "Scoring teams", "Lagged rows", "Notes"],
        [
            ["2016", "dcsaa_2016.txt", "0", "0", "0", "Times without school names"],
            ["2017", "dcsaa_2017.txt", "0", "0", "0", "Junior varsity page"],
            ["2018", "dcsaa_2018.txt", "0", "0", "0", "Junior varsity page"],
            ["2019", "dcsaa_2019.txt", "73", "7", "0", "Lag source"],
            ["2020", "—", "0", "0", "0", "No championship"],
            ["2021", "dcsaa_2021.txt", "71", "8", "6", "Lag from 2019"],
            ["2022", "dcsaa_2022.txt", "85", "12", "8", "Lag and outcome"],
            ["2023", "dcsaa_2023.txt", "93", "14", "11", "Collapsed HY-TEK recovered"],
            ["2024", "dcsaa_2024.txt", "404", "19", "15", "Lag and outcome"],
            ["2025", "dcsaa_2025.txt", "103", "14", "12 holdout", "Temporal holdout"],
        ],
        "Table 1. Championship-year data audit. Venue was Kenilworth Park in every parsed year. Lagged rows are the estimation sample except 2025 (holdout). Public result URLs are listed in the project meet index.",
    )

    add_mixed(doc, [("2.3 Inclusion and feature construction", {"size": 12, "bold": True, "italic": True})], space_after_blank=False)
    add_para(
        doc,
        "Boys’ varsity 5 km finishers were retained. Junior varsity, Varsity B, and middle school races were excluded. Times outside 14–45 min were dropped. Extra finishers were capped at seven per school, matching the usual entry limit rather than a verified official roster. A team entered the regression sample only if five finishers were recorded. Standardization used all parsed varsity boys in that championship, not one observation per school:",
    )
    add_para(doc, "z_i = (t_i − mean_s) / sd_s          (eq1)", center=True, italic=True)
    add_para(
        doc,
        "Higher z-scores indicated slower performance relative to that day’s field. Star z is the first-runner z; depth z is the fifth-runner z; pack gap is depth z minus star z. These quantities partly capture the same school’s competitive strength that later appears as team place. The scientific question is whether last year’s championship form forecasts next year’s place, not whether pack depth has an independent causal effect.",
    )

    add_mixed(doc, [("2.4 Statistical analysis", {"size": 12, "bold": True, "italic": True})], space_after_blank=False)
    add_para(
        doc,
        "Nested ordinary least-squares models of team place used cluster-robust standard errors by school with the finite-sample correction implemented in statsmodels (8). The specifications were: (i) log enrollment and sector, with charter omitted; (ii) controls plus lagged first-runner z; (iii) controls plus lagged fifth-runner z; (iv) controls plus lagged pack gap; (v) both z-scores; and (vi) both z-scores plus WCAC, reported as exploratory. The same nesting was fit for team score. Akaike information criterion (AIC) differences are reported; a small ΔAIC is not treated as decisive (9). Primary holdout metrics were root-mean-square error (RMSE) and mean absolute error versus the controls baseline, for both place and score. Spearman rank correlation was secondary (10). Leave-one-school-out prediction RMSE was the headline robustness check. A wild cluster bootstrap with 399 Rademacher draws supplied a small-sample interval for the depth coefficient. Additional checks included leave-one-season-out RMSE, a permutation test of the depth-model R², and a within-team swap of first- and fifth-runner z that tested the observed in-sample R² gap. Invitational strength-of-schedule was not ingested. Analyses used Python 3 with pandas and statsmodels.",
    )

    add_heading_line(doc, "3. Results")
    add_mixed(doc, [("3.1 Sample", {"size": 12, "bold": True, "italic": True})], space_after_blank=False)
    add_para(
        doc,
        "After lag construction and the five-finisher requirement, 40 team-seasons from 15 schools were available for estimation (2021–2024), and 12 scoring teams were available for 2025 (Figure 1). Files from 2016 listed times without school names. Archives from 2017 and 2018 were junior varsity pages. Prior-year first- and fifth-runner z-scores were strongly correlated (r = 0.84), so the two predictors largely describe the same programs.",
    )
    add_figure(
        doc,
        paths["f1"],
        "Figure 1. Season-by-season inclusion. Gold bars are parsed varsity individual finishers; teal bars are scoring teams; navy bars are lagged rows used in estimation (2021–2024) or holdout (2025).",
    )

    add_mixed(doc, [("3.2 Primary model comparison", {"size": 12, "bold": True, "italic": True})], space_after_blank=False)
    add_para(
        doc,
        "Table 2 reports nested fits with school-clustered 95% confidence intervals. Enrollment and sector alone explained 46.5% of the variance in next-year team place (adjusted R² = 0.42). Private schools placed 6.94 positions better than charter schools (95% CI −10.69 to −3.20). Adding first-runner z raised R² to 0.62 (ΔAIC = 8.1 relative to the fifth-runner model). Adding fifth-runner z raised R² to 0.69 and produced the lowest AIC. The depth coefficient was 3.86 places per standard-deviation slower (cluster 95% CI 2.28 to 5.43; wild-cluster bootstrap 95% CI 2.44 to 5.34). Pack gap alone raised R² only to 0.53 (ΔAIC = 16.7), although the gap coefficient was positive and the interval excluded zero (Figure 2).",
    )
    add_para(
        doc,
        "In the joint z-score model, fifth-runner z remained associated with place (3.13, 95% CI 1.39 to 4.88) while first-runner z did not (1.43, 95% CI −0.59 to 3.46). ΔAIC between the depth-only model and the two-predictor model was 0.82, which is too small to prefer the larger specification (9). Figure 3 shows that both z-scores track next-year place, with a tighter visual association for the fifth runner. Leave-one-school-out prediction RMSE for the depth model was 2.80 places. Nested models of team score ranked specifications the same way: fifth-runner z had the lowest AIC (ΔAIC = 0), versus 15.3 for first-runner z and 25.2 for controls.",
    )
    add_table(
        doc,
        ["Model", "N", "R²", "Adj. R²", "AIC", "ΔAIC", "Focal coef (SE)", "95% CI"],
        [
            ["Controls", "40", "0.465", "0.420", "210.3", "19.7", "Private −6.94 (1.91)", "−10.69, −3.20"],
            ["+ first-runner z", "40", "0.619", "0.575", "198.8", "8.1", "4.32 (1.01)", "2.34, 6.30"],
            ["+ fifth-runner z", "40", "0.689", "0.653", "190.6", "0.0", "3.86 (0.80)", "2.28, 5.43"],
            ["+ pack gap", "40", "0.527", "0.473", "207.4", "16.7", "2.77 (0.93)", "0.95, 4.59"],
            ["Both z-scores", "40", "0.698", "0.654", "191.5", "0.8", "Depth 3.13 (0.89)", "1.39, 4.88"],
            ["+ WCAC (exploratory)", "40", "0.699", "0.644", "193.3", "2.7", "WCAC −0.52 (0.68)", "−1.85, 0.81"],
        ],
        "Table 2. OLS models of next-year DCSAA boys’ team place. School-clustered standard errors (15 schools). ΔAIC is relative to the fifth-runner model.",
    )
    add_figure(
        doc,
        paths["f2"],
        "Figure 2. Prior-year pack gap versus next-year team place in the estimation sample, with an OLS fit and 95% mean confidence band. Lower place is better.",
    )
    add_figure(
        doc,
        paths["f3"],
        "Figure 3. Prior-year first-runner z (left) and fifth-runner z (right) versus next-year team place. Panels share axes. Lower z is faster relative to that year’s field.",
    )

    add_mixed(doc, [("3.3 Holdout and sensitivity", {"size": 12, "bold": True, "italic": True})], space_after_blank=False)
    add_para(
        doc,
        "On 12 scoring teams in 2025, the fifth-runner model ranked schools with Spearman ρ = 0.87 and RMSE = 3.51 places (Figure 4). The controls-only baseline had RMSE = 3.08 and ρ = 0.68. Pack gap had the lowest holdout RMSE (2.53) but a weaker rank correlation (ρ = 0.76) than depth (Table 3). First-runner RMSE was 3.33. Rank order tracked the official results more closely under the depth model, but absolute place error did not improve on enrollment and sector except under the pack-gap specification. The 2025 evaluation is one temporal test, not independent replication.",
    )
    add_table(
        doc,
        ["Model", "Holdout RMSE", "Holdout MAE", "Spearman ρ"],
        [
            ["Controls (baseline)", "3.08", "2.56", "0.68"],
            ["First-runner z", "3.33", "2.43", "0.78"],
            ["Pack gap", "2.53", "2.02", "0.76"],
            ["Fifth-runner z", "3.51", "3.02", "0.87"],
            ["Both z-scores", "3.73", "3.23", "0.87"],
            ["+ WCAC", "3.73", "3.26", "0.92"],
        ],
        "Table 3. 2025 temporal holdout (n = 12). RMSE and MAE are in team places.",
    )
    add_figure(
        doc,
        paths["f4"],
        "Figure 4. Observed versus predicted 2025 team place from the fifth-runner model, with school labels. The dashed line is a perfect numeric match.",
    )
    add_para(
        doc,
        "Leave-one-school-out depth coefficients ranged from 3.25 to 4.61 (mean 3.84). Leave-one-season-out RMSE was 2.02 in 2021, 1.29 in 2022, 2.23 in 2023, and 3.72 in 2024. Permuting fifth-runner z yielded p = 0.005 for the observed depth-model R². Randomly swapping first- and fifth-runner z within teams produced p = 0.29 for the observed in-sample R² advantage of depth, which is the expected consequence of r = 0.84. The exploratory WCAC coefficient was not distinguishable from zero (95% CI −1.85 to 0.81). The girls’ lagged panel contained 18 team-seasons and no 2025 holdout rows with a prior fifth-runner z. Controls explained 66% of girls’ team place; first-runner z raised R² to 0.71 and fifth-runner z raised it to 0.70. The boys’ in-sample ranking of depth over star did not replicate.",
    )

    add_heading_line(doc, "4. Discussion")
    add_para(
        doc,
        "In the DCSAA boys’ estimation sample, prior-year fifth-runner performance added more information about next-year team place than prior-year first-runner performance. Nested R², AIC, the joint model, team-score models, and holdout rank correlation all pointed toward depth. That pattern is consistent with the scoring rule, in which the fifth place counts as much as the first (1). It is not evidence that a school would improve its championship place by slowing its fastest runner or by treating pack work as a substitute for recruiting speed.",
    )
    add_para(
        doc,
        "Three results limit how far that interpretation can be taken. First, first- and fifth-runner z-scores are collinear. The swap test did not distinguish the two predictors, so the in-sample edge for depth is a ranking of two highly overlapping summaries, not a clean isolation of pack structure. Second, holdout RMSE favored pack gap and the controls baseline over the depth-only model. In a 12-team ranking, order can improve while typical numeric error remains several places. Third, the girls’ replication, although small and without a holdout year, did not reproduce a depth advantage. Cluster-robust standard errors with 15 schools are themselves a limited inferential basis (8).",
    )
    add_para(
        doc,
        "A hypothesis, not a demonstrated mechanism, is that the fifth scorer encodes whether a program can still fill a competitive five a year later, whereas first-runner form is noisier from year to year. Enrollment, sector, and WCAC membership are largely known to a coach already; the scientific comparison is whether lagged championship z-scores add information after those terms. WCAC membership did not do so in this panel, and invitational schedules were not scored.",
    )
    add_mixed(doc, [("4.1 Limitations", {"size": 12, "bold": True, "italic": True})], space_after_blank=False)
    add_para(
        doc,
        "The panel is a census of parseable Kenilworth Park championship files, not a sample of U.S. high-school cross-country. Forty lagged team-seasons and 15 schools are a pilot scale. Incomplete teams drop out of the lagged sample, which can select for more stable programs. Enrollment is static. Z-scores depend on who appeared in that year’s field. Ordinary least squares on place is an approximation, even with team score as a co-primary outcome. Associations are not causal.",
    )
    add_mixed(doc, [("4.2 Future work", {"size": 12, "bold": True, "italic": True})], space_after_blank=False)
    add_para(
        doc,
        "Recovered 2017–2018 varsity individual lists would enlarge the lagged panel without leaving the district. A larger girls’ sample with a reserved holdout year would test whether the first-versus-fifth ranking is sex-specific. Invitational team–meet rows would measure schedule strength more directly than conference membership (4, 5). Multi-district championships on comparable courses would be required before treating fifth-runner z as a general forecasting tool.",
    )

    add_heading_line(doc, "5. Conclusion")
    add_para(
        doc,
        "In this small DCSAA boys’ panel, prior-year fifth-runner within-meet z-score predicted next-year team place more closely in-sample than prior-year first-runner z-score, including in a joint model, with leave-one-school-out RMSE of 2.80 places. A frozen depth model reproduced the 2025 rank ordering (Spearman ρ = 0.87), while RMSE remained about 3.5 places and did not beat a controls-only baseline. Pack gap had the lowest holdout RMSE. The girls’ replication did not show the same depth advantage. The result is consistent with equal weighting of the first five places in the scoring rule, but it does not test that rule and does not support a causal claim about pack depth.",
    )

    add_heading_line(doc, "Data availability")
    add_para(
        doc,
        "Cleaned team-season and individual files, the season audit, the exclusion log, the school alias table, and model output are provided with the analysis code. The pipeline that constructs the panel and the script that compiles this manuscript are included in the same repository.",
    )
    add_heading_line(doc, "Acknowledgements")
    add_para(doc, "Public championship listings from DCSAA, MileSplit, and M&D Timing were used. No external funding was received.")
    add_heading_line(doc, "Abbreviations")
    add_para(
        doc,
        "AIC, Akaike information criterion; DCSAA, District of Columbia State Athletic Association; MAE, mean absolute error; NFHS, National Federation of State High School Associations; OLS, ordinary least squares; RMSE, root-mean-square error; WCAC, Washington Catholic Athletic Conference.",
    )
    add_heading_line(doc, "References")
    for ref in [
        "(1) National Federation of State High School Associations. 2025–26 NFHS Track and Field and Cross Country Rules Book. NFHS, Indianapolis, IN, USA, 2025.",
        "(2) Stefani RT. Improved least squares football, basketball, and soccer ratings. IEEE Trans Syst Man Cybern, 10: 116–123, 1980. https://doi.org/10.1109/TSMC.1980.4308442",
        "(3) Harville DA, Smith MH. The home-court advantage: How large is it, and does it vary from team to team? Am Stat, 48: 22–28, 1994. https://doi.org/10.1080/00031305.1994.10476013",
        "(4) Kvam P, Sokol JS. A logistic regression/Markov chain model for NCAA basketball. Nav Res Logist, 53: 788–803, 2006. https://doi.org/10.1002/nav.20165",
        "(5) Pasteur RD, Janning MC. Monte Carlo simulation for high school football playoff seed projection. J Quant Anal Sports, 7(2): 11, 2011. https://doi.org/10.2202/1559-0410.1330",
        "(6) Johnson JE, Manwell AK, Scott BF. An examination of competitive balance within interscholastic football. J Amateur Sport, 5(1): 21–49, 2019. https://doi.org/10.17161/jas.v5i1.6708",
        "(7) District of Columbia State Athletic Association. 2025 Cross Country Bulletin. DCSAA, Washington, DC, USA, 2025. https://www.dcsaasports.com/boys-cross-country/",
        "(8) Cameron AC, Miller DL. A practitioner’s guide to cluster-robust inference. J Hum Resour, 50: 317–372, 2015. https://doi.org/10.3368/jhr.50.2.317",
        "(9) Akaike H. A new look at the statistical model identification. IEEE Trans Automat Contr, 19: 716–723, 1974. https://doi.org/10.1109/TAC.1974.1100705",
        "(10) Spearman C. The proof and measurement of association between two things. Am J Psychol, 15: 72–101, 1904. https://doi.org/10.2307/1412159",
    ]:
        add_para(doc, ref, size=12)

    doc.save(DOCX)
    return DOCX


def main():
    paths = make_figures()
    out = build_docx(paths)
    print("wrote", out)
    for k, v in paths.items():
        print(k, v)


if __name__ == "__main__":
    main()
