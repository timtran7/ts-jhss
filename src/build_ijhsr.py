"""IJHSR research-article Word file, Arial, ACS superscripts, PNG figures."""

from __future__ import annotations

import re
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import statsmodels.formula.api as smf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.build_manuscript import make_figures

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "paper" / "figures"
PNG = ROOT / "paper" / "ijhsr_png"
DOCX = ROOT / "paper" / "LastName_FirstName.docx"
PDF = ROOT / "paper" / "LastName_FirstName.pdf"
FONT = "Arial"


def _set_run(run, size=10, bold=False, italic=False, superscript=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.superscript = superscript
    if color:
        run.font.color.rgb = RGBColor(*color)


def _pf(p, after=0, before=0, single=True):
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0 if single else 1.15
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE if single else WD_LINE_SPACING.MULTIPLE


def heading(doc, text: str):
    p = doc.add_paragraph()
    _pf(p, after=8, before=8)
    r = p.add_run(text)
    _set_run(r, size=14, bold=True)
    return p


def subhead(doc, text: str):
    p = doc.add_paragraph()
    _pf(p, after=8, before=6)
    r = p.add_run(text)
    _set_run(r, size=12, bold=True, italic=True)
    return p


def body(doc, text: str, *, size=10, bold=False, italic=False, center=False, after=8):
    """Body text. Use |12| or |3-5| after punctuation for ACS superscripts."""
    p = doc.add_paragraph()
    _pf(p, after=after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = re.split(r"(\|\d+(?:[-–,]\d+)*\|)", text)
    for part in parts:
        if not part:
            continue
        m = re.fullmatch(r"\|(\d+(?:[-–,]\d+)*)\|", part)
        if m:
            r = p.add_run(m.group(1).replace(",", ", "))
            _set_run(r, size=size, superscript=True)
        else:
            r = p.add_run(part)
            _set_run(r, size=size, bold=bold, italic=italic)
    return p


def shade_header(cell, hex_color="1B365D"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, caption: str):
    cap = doc.add_paragraph()
    _pf(cap, after=8)
    r = cap.add_run(caption)
    _set_run(r, size=11, italic=False)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        _pf(p, after=0)
        run = p.add_run(h)
        _set_run(run, size=9, bold=True, color=(255, 255, 255))
        shade_header(cell)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _pf(p, after=0)
            run = p.add_run(val)
            _set_run(run, size=8)
    spacer = doc.add_paragraph()
    _pf(spacer, after=8)


def add_figure(doc, path: Path, caption: str, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _pf(p, after=4)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    _pf(cap, after=8)
    r = cap.add_run(caption)
    _set_run(r, size=11)


def export_png(jpgs: dict[str, Path]) -> dict[str, Path]:
    PNG.mkdir(parents=True, exist_ok=True)
    names = {
        "f1": "Figure1.png",
        "f2": "Figure2.png",
        "f3": "Figure3.png",
        "f4": "Figure4.png",
    }
    out = {}
    for key, src in jpgs.items():
        dest = PNG / names[key]
        # Re-save via matplotlib-independent copy: Pillow if present, else copy jpg bytes with png name via PIL
        try:
            from PIL import Image

            Image.open(src).save(dest, "PNG")
        except Exception:
            dest = dest.with_suffix(".png")
            dest.write_bytes(Path(src).read_bytes())
        out[key] = dest
    return out


def build_docx(paths: dict[str, Path]) -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.first_line_indent = Inches(0)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _pf(t, after=8)
    r = t.add_run("Star Power or Pack Depth? Predicting DCSAA Boys Cross-Country Team Place from Last Year’s Results")
    _set_run(r, size=16, bold=True)

    body(doc, "[Author Name] (student)", center=True, after=0)
    body(doc, "[High School], [Street Address], [City], [State ZIP], USA", center=True, after=0)
    body(doc, "* Corresponding author email: [personal@email.com]", center=True, after=16)

    heading(doc, "Abstract")
    body(
        doc,
        "In high school cross country, a team’s score is the sum of its first five finishing places. This study asked whether last year’s fifth runner predicts next year’s team place better than last year’s first runner at the District of Columbia State Athletic Association (DCSAA) boys varsity championship. Public Kenilworth Park results were used. Each runner’s time was compared with the rest of that year’s field. Models also accounted for school size and whether the school was private, public, or charter. The 2025 championship was left out of model fitting and used as a check. Last year’s fifth runner was the stronger predictor in the years used to fit the models. In 2025 it ranked the teams in close to the official order, but the size of the place errors was not better than using school size and type alone. These findings are for one district and should not be read as proof that pack depth causes a better team place.",
        after=8,
    )

    heading(doc, "Keywords")
    for kw in [
        "Computer Science and Software Engineering",
        "Data Science",
        "Sports analytics",
        "Cross-country running",
        "Predictive modeling",
    ]:
        body(doc, kw, after=0)
    body(doc, "", after=8)

    heading(doc, "Introduction")
    body(
        doc,
        "In high school cross country, a team’s score is the sum of the first five finishers’ places.|1| A school can have the fastest runner in the race and still lose if its fifth scorer is far behind. Coaches often describe this as star power versus pack depth. There is not much published work that tests which one actually predicts next year’s championship.",
    )
    body(
        doc,
        "Using the same race to predict the same team’s place would not be a fair test. The fifth runner’s place is already part of that day’s score. A more useful test uses last year’s championship to predict this year’s place, because those results are already known before the next race.",
    )
    body(
        doc,
        "Sports rating papers have used least squares to adjust for opponent strength.|2,3| Other work has used schedule measures in college basketball and simulation for high school football playoffs.|4,5| Public and private schools also differ in many high school sports.|6| Those studies are about rankings, seeding, or competitive balance. They do not compare last year’s first runner and last year’s fifth runner as predictors of a high school cross country championship.",
    )
    body(
        doc,
        "DCSAA holds one boys varsity championship instead of separate class races.|7| The Kenilworth Park course has been 5 km. Teams usually enter up to seven runners and need five finishers to score.|7| Because the course is the same each year, times can be turned into z-scores within that meet. The goal of this research is to test whether last year’s fifth-runner z-score predicts next year’s DCSAA boys team place better than last year’s first-runner z-score after enrollment and sector. Washington Catholic Athletic Conference (WCAC) membership is included as an extra check for schedule strength. This paper does not claim that changing fifth-runner times would cause a change in team place.",
    )

    heading(doc, "Methods")
    subhead(doc, "Design and Outcomes")
    body(
        doc,
        "Each row is one school in one season. The main outcome is team place at the DCSAA boys championship. Team score (the sum of the first five places) is also modeled because that is how the meet is scored. Place is a ranking, so ordinary least squares is only an approximation. Standard errors are clustered by school.|8| A wild cluster bootstrap with 399 draws is used for the fifth-runner coefficient. Predictors come only from the previous Kenilworth championship that could be parsed. There was no 2020 race, so 2021 is linked to 2019. After the 2023 file was recovered, 2024 is linked to 2023. The 2025 season was not used to pick the models.",
    )
    subhead(doc, "Data Sources")
    body(
        doc,
        "Meet pages came from the public MileSplit championship index and, for 2025, M&D Timing (Table 1).|7,9| Local copies of the pages are stored with the code. If a posted team-score table and a reconstructed table both existed, the posted table was used. Individual times were used for first- and fifth-runner z-scores and to rebuild scores when no team table was posted. School names were matched with an alias list that also has sector, WCAC status, and a static grades 9–12 enrollment estimate. Names that could not be matched were dropped. Analyses used Python 3 with pandas, NumPy, and statsmodels.|10–13|",
    )
    add_table(
        doc,
        ["Season", "File", "Finishers parsed", "Scoring teams", "Lagged rows", "Notes"],
        [
            ["2016", "dcsaa_2016.txt", "0", "0", "0", "Times without school names"],
            ["2017", "dcsaa_2017.txt", "0", "0", "0", "Junior varsity page"],
            ["2018", "dcsaa_2018.txt", "0", "0", "0", "Junior varsity page"],
            ["2019", "dcsaa_2019.txt", "73", "7", "0", "Lag source"],
            ["2020", "—", "0", "0", "0", "No championship"],
            ["2021", "dcsaa_2021.txt", "71", "8", "6", "Lag from 2019"],
            ["2022", "dcsaa_2022.txt", "85", "12", "8", "Lag and outcome"],
            ["2023", "dcsaa_2023.txt", "93", "14", "11", "Collapsed HY-TEK recovered"],
            ["2024", "dcsaa_2024.txt", "404", "19", "15", "Lag and outcome"],
            ["2025", "dcsaa_2025.txt", "103", "14", "12 holdout", "Temporal holdout"],
        ],
        "Table 1. Championship-year data audit. All parsed years were at Kenilworth Park. Lagged rows are the fit sample except 2025 (holdout).",
    )
    subhead(doc, "Inclusion and Features")
    body(
        doc,
        "Only boys varsity 5 km finishers were kept. Junior varsity, Varsity B, and middle school races were dropped. Times outside 14–45 min were dropped. Extra finishers were capped at seven per school. A team needed five recorded finishers to enter the models. The z-score was z_i = (t_i − mean_s) / sd_s, using all parsed varsity boys that day. A higher z-score means a slower time relative to that field. Star z is the first runner. Depth z is the fifth runner. Pack gap is depth z minus star z.",
    )
    subhead(doc, "Statistical Analysis")
    body(
        doc,
        "Nested OLS models of team place were: (i) log enrollment and sector (charter omitted); (ii) controls plus last year’s first-runner z; (iii) controls plus last year’s fifth-runner z; (iv) controls plus last year’s pack gap; (v) both z-scores; (vi) both z-scores plus WCAC. Model (vi) is exploratory. The same models were fit for team score. AIC differences are reported; a small ΔAIC is not treated as a clear winner.|14| The 2025 season is one holdout year. RMSE and MAE versus the controls model are the main holdout metrics for place and score. Spearman correlation is secondary.|15| Leave-one-school-out RMSE is the main extra check. Other checks were leave-one-season-out RMSE, a permutation test of the fifth-runner model R², and a swap of first- and fifth-runner z within teams. Invitational meet results were not used.",
    )

    heading(doc, "Results and Discussion")
    subhead(doc, "Sample")
    body(
        doc,
        "After requiring a lag and five finishers, 40 team-seasons from 15 schools were in the fit sample (2021–2024). Twelve scoring teams were in 2025 (Figure 1). Last year’s first- and fifth-runner z-scores were highly correlated (r = 0.84).",
    )
    add_figure(
        doc,
        paths["f1"],
        "Figure 1. Season-by-season inclusion of parsed varsity finishers, scoring teams, and lagged rows. Gold bars are individual finishers; teal bars are scoring teams; navy bars are lagged fit or holdout rows. This graph was created using Python 3 (matplotlib 3).|13|",
        width=6.5,
    )
    subhead(doc, "Primary Model Comparison")
    body(
        doc,
        "Table 2 shows the nested OLS fits. Controls explained 46.5% of the variance in team place (adjusted R² = 0.42). Adding first-runner z raised R² to 0.62 (ΔAIC = 8.1 versus the fifth-runner model). Adding fifth-runner z raised R² to 0.69 and had the lowest AIC. The fifth-runner coefficient was 3.86 places per SD slower (cluster 95% CI 2.28 to 5.43; wild cluster bootstrap 95% CI 2.44 to 5.34). Pack gap alone only reached R² = 0.53 (ΔAIC = 16.7). In the joint model, fifth-runner z was still related to place (3.13, 95% CI 1.39 to 4.88) and first-runner z was not (1.43, 95% CI −0.59 to 3.46). Leave-one-school-out RMSE for the fifth-runner model was 2.80 places (Figure 2 and Figure 3). Team-score models ranked the same way. Fifth-runner z had the lowest AIC.",
    )
    add_table(
        doc,
        ["Model", "N", "R²", "Adj. R²", "AIC", "ΔAIC", "Focal coef (SE)", "95% CI"],
        [
            ["Controls", "40", "0.465", "0.420", "210.3", "19.7", "Private −6.94 (1.91)", "−10.69, −3.20"],
            ["+ first-runner z", "40", "0.619", "0.575", "198.8", "8.1", "4.32 (1.01)", "2.34, 6.30"],
            ["+ fifth-runner z", "40", "0.689", "0.653", "190.6", "0.0", "3.86 (0.80)", "2.28, 5.43"],
            ["+ pack gap", "40", "0.527", "0.473", "207.4", "16.7", "2.77 (0.93)", "0.95, 4.59"],
            ["Both z-scores", "40", "0.698", "0.654", "191.5", "0.8", "Depth 3.13 (0.89)", "1.39, 4.88"],
            ["+ WCAC (exploratory)", "40", "0.699", "0.644", "193.3", "2.7", "WCAC −0.52 (0.68)", "−1.85, 0.81"],
        ],
        "Table 2. OLS models of next-year DCSAA boys team place with school-clustered standard errors (15 schools). ΔAIC is relative to the fifth-runner model. Statistical analysis was performed in Python 3 using statsmodels.|10,12|",
    )
    add_figure(
        doc,
        paths["f2"],
        "Figure 2. Last year’s pack gap versus next year’s team place in the fit sample, with an OLS line and 95% mean confidence band. Lower place is better. This graph was created using Python 3 (matplotlib 3).|13|",
        width=5.5,
    )
    add_figure(
        doc,
        paths["f3"],
        "Figure 3. Last year’s first-runner z (left) and fifth-runner z (right) versus next year’s team place. Panels share axes. Lower z is faster relative to that year’s field. This graph was created using Python 3 (matplotlib 3).|13|",
        width=6.8,
    )
    subhead(doc, "Holdout and Sensitivity")
    body(
        doc,
        "On 12 scoring teams in 2025, the fifth-runner model had RMSE = 3.51 and Spearman ρ = 0.87 (Figure 4). The controls model had RMSE = 3.08 and ρ = 0.68. Pack gap had the lowest holdout RMSE (2.53) but a lower rank correlation (ρ = 0.76) than fifth-runner z (Table 3). The fifth-runner model ranked the teams closer to the official order, but the typical place error was not better than enrollment and sector except for pack gap.",
    )
    add_table(
        doc,
        ["Model", "Holdout RMSE", "Holdout MAE", "Spearman ρ"],
        [
            ["Controls (baseline)", "3.08", "2.56", "0.68"],
            ["First-runner z", "3.33", "2.43", "0.78"],
            ["Pack gap", "2.53", "2.02", "0.76"],
            ["Fifth-runner z", "3.51", "3.02", "0.87"],
            ["Both z-scores", "3.73", "3.23", "0.87"],
            ["+ WCAC", "3.73", "3.26", "0.92"],
        ],
        "Table 3. 2025 temporal holdout (n = 12). RMSE and MAE are in team places. Metrics were computed in Python 3 using pandas and NumPy.|10,11|",
    )
    add_figure(
        doc,
        paths["f4"],
        "Figure 4. Observed versus predicted 2025 team place from the fifth-runner model, with school labels. The dashed line is a perfect numeric match. This graph was created using Python 3 (matplotlib 3).|13|",
        width=5.5,
    )
    body(
        doc,
        "Leave-one-school-out fifth-runner coefficients ranged from 3.25 to 4.61 (mean 3.84). Leave-one-season-out RMSE was 2.02 (2021), 1.29 (2022), 2.23 (2023), and 3.72 (2024). A permutation test of fifth-runner z gave p = 0.005 for the observed R². Swapping first- and fifth-runner z within teams gave p = 0.29 for the in-sample R² gap, which matches the high correlation (r = 0.84). The WCAC coefficient was not significant (95% CI −1.85 to 0.81).",
    )
    body(
        doc,
        "For the boys sample, last year’s fifth runner added more information about next year’s team place than last year’s first runner. R², AIC, the joint model, and holdout rank correlation all pointed that way. Holdout RMSE did not. In a 12-team race, the order can look better even when the typical error is still a few places. One possible reason, which this study does not prove, is that each of the first five places counts the same, so the fifth scorer may show whether a school can still fill a scoring five the next year.|1| Fifteen schools is a small number of clusters for clustered standard errors.|8| 2025 is only one holdout year.",
    )
    subhead(doc, "Limitations")
    body(
        doc,
        "This is one district and one course. Forty team-seasons and 15 schools is a small sample. Enrollment does not change by year. Teams that did not finish five runners drop out of the lag sample. Invitationals were not scored. OLS on place is an approximation. The results are associations, not cause and effect. The same lag setup on girls varsity files (18 team-seasons, no 2025 holdout) did not rank fifth runner above first runner, but that sample is too small to treat as a second study. For boys in this sample, watching the likely fifth scorer may help a forecast more than watching only the fastest runner, but 2025 RMSE was better for pack gap and for controls than for the fifth-runner-only model.",
    )

    heading(doc, "Conclusion")
    body(
        doc,
        "In this DCSAA boys sample, last year’s fifth-runner z-score predicted next year’s team place better in-sample than last year’s first-runner z-score. Leave-one-school-out RMSE was 2.80 places. A frozen fifth-runner model ranked the 2025 teams with Spearman ρ = 0.87. RMSE was about 3.5 places and did not beat the controls model. Pack gap had the lowest holdout RMSE. More districts would be needed before using pack depth as a general forecasting tool.",
    )

    heading(doc, "AI Acknowledgement")
    body(
        doc,
        "During the preparation of this work, the author used Cursor to help organize Python code for parsing results, fitting models, making figures, and formatting the Word file. After using this tool, the author chose the research question and the models, checked the output, and edited the text. The author takes full responsibility for the paper. No AI tool was used to invent data, citations, or conclusions. Figures were made in Python (matplotlib) and are not AI-generated images.",
    )

    heading(doc, "Acknowledgements")
    body(
        doc,
        "Public championship listings from DCSAA, MileSplit, and M&D Timing were used. No external funding was received. [Mentor Name] provided guidance on the publication process.",
    )

    heading(doc, "References")
    refs = [
        "1. National Federation of State High School Associations. 2025–26 NFHS Track and Field and Cross Country Rules Book; NFHS: Indianapolis, IN, 2025.",
        "2. Stefani, R. T. Improved least squares football, basketball, and soccer ratings. IEEE Trans. Syst. Man Cybern. 1980, 10, 116–123. DOI: 10.1109/TSMC.1980.4308442",
        "3. Harville, D. A.; Smith, M. H. The home-court advantage: How large is it, and does it vary from team to team? Am. Stat. 1994, 48, 22–28. DOI: 10.1080/00031305.1994.10476013",
        "4. Kvam, P.; Sokol, J. S. A logistic regression/Markov chain model for NCAA basketball. Nav. Res. Logist. 2006, 53, 788–803. DOI: 10.1002/nav.20165",
        "5. Pasteur, R. D.; Janning, M. C. Monte Carlo simulation for high school football playoff seed projection. J. Quant. Anal. Sports 2011, 7 (2), 11. DOI: 10.2202/1559-0410.1330",
        "6. Johnson, J. E.; Manwell, A. K.; Scott, B. F. An examination of competitive balance within interscholastic football. J. Amateur Sport 2019, 5 (1), 21–49. DOI: 10.17161/jas.v5i1.6708",
        "7. District of Columbia State Athletic Association. 2025 Cross Country Bulletin. https://www.dcsaasports.com/boys-cross-country/ (accessed 2026-08-15).",
        "8. Cameron, A. C.; Miller, D. L. A practitioner’s guide to cluster-robust inference. J. Hum. Resour. 2015, 50, 317–372. DOI: 10.3368/jhr.50.2.317",
        "9. MileSplit. DCSAA Cross Country Championships results index. https://dc.milesplit.com/ (accessed 2026-08-15).",
        "10. Van Rossum, G.; Drake, F. L. Python 3 Reference Manual; Python Software Foundation: Wilmington, DE. https://www.python.org/ (accessed 2026-08-15).",
        "11. McKinney, W. Data structures for statistical computing in Python. Proc. 9th Python Sci. Conf. 2010, 56–61. DOI: 10.25080/Majora-92bf1922-00a",
        "12. Seabold, S.; Perktold, J. Statsmodels: Econometric and statistical modeling with Python. Proc. 9th Python Sci. Conf. 2010, 92–96. DOI: 10.25080/Majora-92bf1922-011",
        "13. Hunter, J. D. Matplotlib: A 2D graphics environment. Comput. Sci. Eng. 2007, 9, 90–95. DOI: 10.1109/MCSE.2007.55",
        "14. Akaike, H. A new look at the statistical model identification. IEEE Trans. Autom. Control 1974, 19, 716–723. DOI: 10.1109/TAC.1974.1100705",
        "15. Spearman, C. The proof and measurement of association between two things. Am. J. Psychol. 1904, 15, 72–101. DOI: 10.2307/1412159",
    ]
    for ref in refs:
        body(doc, ref, after=4)

    heading(doc, "Authors")
    body(
        doc,
        "[Author Name] is a high school student. This project uses public championship results to compare last year’s first runner and fifth runner as predictors of next year’s DCSAA boys team place. Intended college and major: [to be completed].",
        after=8,
    )

    doc.save(DOCX)
    return DOCX


def export_pdf(docx_path: Path, pdf_path: Path) -> Path:
    docx_path = docx_path.resolve()
    pdf_path = pdf_path.resolve()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return pdf_path
    except Exception:
        pass
    import subprocess

    ps = (
        f"$word = New-Object -ComObject Word.Application; $word.Visible = $false; "
        f"$doc = $word.Documents.Open('{docx_path}'); "
        f"$pdf = '{pdf_path}'; $doc.SaveAs([ref]$pdf, [ref]17); "
        f"$doc.Close($false); $word.Quit()"
    )
    subprocess.run(["powershell", "-NoProfile", "-Command", ps], check=True)
    return pdf_path


def main():
    jpgs = make_figures()
    pngs = export_png(jpgs)
    out = build_docx(jpgs)
    print("wrote", out)
    pdf = export_pdf(out, PDF)
    print("wrote", pdf)
    for k, v in pngs.items():
        print(k, v)


if __name__ == "__main__":
    main()
