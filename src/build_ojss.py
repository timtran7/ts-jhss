"""Oxford Journal of Student Scholarship original-research Word/PDF file."""

from __future__ import annotations

from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.figures import make_figures
from src.report import _pct, _r, export_pdf, load_report

ROOT = Path(__file__).resolve().parents[1]
PNG = ROOT / "paper" / "ojss_png"
DOCX = ROOT / "paper" / "Tran_Timothy.docx"
PDF = ROOT / "paper" / "Tran_Timothy.pdf"
FONT = "Times New Roman"
REPO = "https://github.com/timtran7/ts-jhss"
AUTHOR = "Timothy Tran"
SCHOOL = "St. Albans School"
GRADE = "12th grade"
CITY = "Washington, DC, USA"
EMAIL = "timotran27@gmail.com"
TITLE = (
    "Prior First- and Fifth-Runner Times as Predictors of Championship "
    "Team Place in District of Columbia High School Cross Country"
)


def _set_run(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _pf(p, *, after=0, before=0, first_line=0.5, space=2.0, align=None):
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.first_line_indent = Inches(first_line)
    p.paragraph_format.line_spacing = space
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if align is not None:
        p.alignment = align


def body(doc, text: str, *, first_line=0.5):
    p = doc.add_paragraph()
    _pf(p, first_line=first_line)
    r = p.add_run(text)
    _set_run(r)
    return p


def heading(doc, text: str):
    p = doc.add_paragraph()
    _pf(p, first_line=0, before=12, after=0)
    r = p.add_run(text)
    _set_run(r, size=12, bold=True)
    return p


def subhead(doc, text: str):
    p = doc.add_paragraph()
    _pf(p, first_line=0, before=6, after=0)
    r = p.add_run(text)
    _set_run(r, size=12, bold=True, italic=True)
    return p


def caption(doc, text: str):
    p = doc.add_paragraph()
    _pf(p, first_line=0, after=6)
    r = p.add_run(text)
    _set_run(r, size=12, italic=True)
    return p


def add_page_number(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _pf(p, first_line=0, space=1.0)
    run = p.add_run()
    _set_run(run, size=12)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def shade_header(cell, hex_color="E8E8E8"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, title: str):
    caption(doc, title)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        _pf(p, first_line=0, space=1.0)
        run = p.add_run(h)
        _set_run(run, size=12, bold=True)
        shade_header(cell)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _pf(p, first_line=0, space=1.0)
            run = p.add_run(str(val))
            _set_run(run, size=12)
    spacer = doc.add_paragraph()
    _pf(spacer, first_line=0)


def add_figure(doc, path: Path, title: str, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _pf(p, first_line=0, after=0)
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, title)


def export_png(jpgs: dict[str, Path]) -> dict[str, Path]:
    PNG.mkdir(parents=True, exist_ok=True)
    names = {"f1": "Figure1.png", "f4": "Figure2.png", "f2": "Figure3.png", "f3": "Figure4.png"}
    out = {}
    for key, src in jpgs.items():
        dest = PNG / names[key]
        Image.open(src).save(dest, "PNG", dpi=(300, 300))
        out[key] = dest
    return out


def hanging(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(text)
    _set_run(r)
    return p


def build_docx(paths: dict[str, Path]) -> Path:
    rep = load_report()
    m = rep["m"]
    g = rep["g"]
    n = rep["n"]
    n_schools = rep["n_schools"]
    corr = rep["corr"]
    hold_n = rep["hold_n"]
    recon = m.get("score_reconciliation", {})
    roll = m.get("rolling_origin", {})
    vif = m.get("vif_joint", {})
    pcorr = m.get("partial_corr_star_depth")
    rmse_diff = m.get("holdout_rmse_diff", {})
    soft = m.get("software", {})
    cov = m.get("holdout_coverage", {})
    scoring_hold = cov.get("scoring_teams", hold_n)
    pc, ps, pd_ = m["controls"], m["star"], m["depth"]
    pg, pb, pw = m["pack_gap"], m["both"], m["wcac"]
    sc, ss, sd, sb = m["score_controls"], m["score_star"], m["score_depth"], m["score_both"]
    wild = m["wild_cluster_depth"]
    ologit = m["ordered_logit_delta_aic"]
    loo_rmse = m["loo_school_place_rmse"]["rmse"]
    hd, hc, hs = m["holdout_by_model"]["depth"], m["holdout_by_model"]["controls"], m["holdout_by_model"]["star"]
    hg, hb, hw = m["holdout_by_model"]["pack_gap"], m["holdout_by_model"]["both"], m["holdout_by_model"]["wcac"]
    hsc = m["holdout_score_by_model"]
    loo_c = m["loo_school_depth_coef"]
    perm = m["permutation_depth_r2"]
    swap = m["permutation_depth_minus_star_r2"]
    gc, gs, gd, gh = g["controls"], g["star"], g["depth"], g["holdout_by_model"]

    dms = rmse_diff.get("depth_minus_star", {})
    dmc = rmse_diff.get("depth_minus_controls", {})

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.different_first_page_header_footer = True
        add_page_number(section)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Title page (OJSS)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _pf(t, first_line=0, before=72)
    r = t.add_run(TITLE)
    _set_run(r, size=18, bold=True)

    for line in [
        AUTHOR,
        SCHOOL,
        GRADE,
        CITY,
        EMAIL,
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _pf(p, first_line=0)
        rr = p.add_run(line)
        _set_run(rr, size=12)

    doc.add_page_break()

    heading(doc, "Abstract")
    p = doc.add_paragraph()
    _pf(p, first_line=0)
    r = p.add_run(
        "High school cross country awards team titles by summing the finishing places of five scorers, "
        "which raises the question of whether next championship team place is better anticipated by a "
        "program’s fastest runner or by its fifth scorer. Prior work on packing looks at how tightly "
        "teammates cluster during a race, while sports-forecasting studies rate teams from repeated "
        "game outcomes. Neither literature compares lagged first- versus fifth-runner summaries under "
        "a five-scorer rule. This paper tests that comparison using public varsity championship results "
        "from the District of Columbia State Athletic Association, all contested at five kilometers on "
        "a single course. Next championship team place is the primary outcome. Models adjust for school "
        "enrollment and sector and are evaluated on a later championship that was not used to choose "
        "the specification. Prior fifth-runner relative time was more strongly associated with "
        "subsequent boys’ team place than prior first-runner relative time, and it reduced holdout "
        "error relative to school characteristics alone. Models that included both runners or "
        "conference membership had even lower error on that later meet, though the two runner "
        "summaries were highly collinear. An exploratory girls’ panel did not reproduce the same "
        "in-sample ranking. The result is a local ranking of public summaries, not a causal account of "
        "pack depth or a forecasting model that should be expected to travel."
    )
    _set_run(r)

    heading(doc, "Introduction")
    body(
        doc,
        "Team scoring in high school cross country is not the same as individual placing. Under "
        "National Federation rules, team place is determined by summing the finishing places of a "
        "school’s first five runners, and sixth and seventh finishers do not add to their own team’s "
        "score even though they displace opposing scorers who finish behind them (National Federation "
        "of State High School Associations, 2025). A program can have the meet’s fastest individual "
        "and still lose the team title if its fifth scorer is far back in the field.",
    )
    body(
        doc,
        "That scoring rule leaves two competing summaries of prior performance. One is the lead "
        "runner, whose time usually corresponds to a high place and is often treated as the program’s "
        "ceiling. The other is the fifth scorer, who sits at the margin of the scoring five. The "
        "question here is which of those lagged championship summaries, if either, better predicts the "
        "next team result after adjustment for school size and sector.",
    )
    body(
        doc,
        "The packing literature is related, but it asks a different question. Falk, Larson, and "
        "DeBeliso (2022) measured how often NCAA women’s teammates ran within one second of one "
        "another at race checkpoints and found that greater packing was associated with faster "
        "combined team times in the same race. If that within-race construct were interchangeable with "
        "a lagged fifth-runner time, the spread between a team’s first and fifth times (“pack gap”) "
        "should be the more informative predictor of next championship place. This paper tests that "
        "implication directly. Finley and Fountain (2023) likewise remain inside a single championship, "
        "classifying NCAA teams by opening pace and subsequent movement through the field relative to "
        "pre-meet rankings. Hanley (2015, 2016) documented packing and pacing in world-level half "
        "marathons and championship marathons, including sex differences in elite fields that are "
        "deeper and more homogeneous than a District of Columbia high-school championship. Those papers "
        "describe what happens during a race. They do not compare last championship’s first runner "
        "with last championship’s fifth runner as forecasts of next team place.",
    )
    body(
        doc,
        "Sports-forecasting research is more useful here as an evaluation standard than as a model to "
        "copy. Stefani (1980) constructed least-squares ratings from game results. Kvam and Sokol "
        "(2006) combined logistic regression with a Markov chain using a season of NCAA basketball "
        "outcomes. Pasteur and Janning (2011) simulated remaining high-school football schedules to "
        "project playoff seeds. Those designs rely on repeated contests against identified opponents. "
        "A District of Columbia cross country championship is a single annual race scored by five "
        "internal places, so applying a full rating system to this file would overstate the available "
        "information. What those studies motivate is out-of-sample evaluation: a predictor should be "
        "judged on meets that were not used to select it.",
    )
    body(
        doc,
        "School resources are also a plausible confounder. Johnson, Manwell, and Scott (2019) found "
        "that private schools win a disproportionate share of high-school football state titles and "
        "that the public–private distinction outweighed several other structural variables. The "
        "District of Columbia championship is an open meet mixing independent, public, and charter "
        "schools, several of which also compete in the Washington Catholic Athletic Conference. "
        "Comparing first- and fifth-runner times without enrollment and sector would risk attributing "
        "to “depth” what is partly a school-type difference, so those controls belong in the "
        "confirmatory specification.",
    )
    body(
        doc,
        "The District of Columbia State Athletic Association contests a single open championship. In "
        "the parseable Kenilworth Park years, the race was held at five kilometers with a usual "
        "seven-runner cap (District of Columbia State Athletic Association, 2025). The fixed venue "
        "makes within-meet standardization possible without converting times across courses. The "
        "contribution of this paper is a site-specific ranking of two lagged public summaries as "
        "predictors of next boys’ varsity team place among schools with an eligible prior championship, "
        "after enrollment and sector adjustment. Accuracy is defined as lower error on a reserved later "
        "championship. The paper does not propose a general forecasting system, and it does not treat "
        "fifth-runner relative time as a causal measure of pack depth.",
    )

    heading(doc, "Methods")
    subhead(doc, "Data Sources and Sample")
    body(
        doc,
        "This paper uses public championship archives rather than invitational results so that every "
        "predictor is known before the outcome race and is measured on the same course. Meet listings "
        "were taken from the MileSplit championship index and, for 2025, from M&D Timing (District of "
        "Columbia State Athletic Association, 2025; MileSplit, 2026; M&D Timing, 2025). Local copies of "
        "result pages are stored with the analysis code. When a posted team-score table and a reconstructed "
        "table both existed, the posted table was treated as authoritative.",
    )
    if recon.get("n_overlap"):
        body(
            doc,
            f"Where posted and reconstructed scores overlapped (n = {recon['n_overlap']}), exact place "
            f"agreement was {_r(100 * recon.get('place_exact_match_rate', 0), 1)}% and exact score agreement "
            f"was {_r(100 * recon.get('score_exact_match_rate', 0), 1)}%. Reconstructed scores were used only "
            f"when a posted team table was missing.",
        )
    body(
        doc,
        "School names were mapped through a directory that records sector (private, public, or charter), "
        "Washington Catholic Athletic Conference membership, and grades 9–12 enrollment. Public-school "
        "enrollment was taken from the official fall audit for the school year containing the championship "
        "(District of Columbia Public Schools, 2026). Private and charter enrollment used a static directory "
        "estimate when year-specific audits were unavailable, and unmatched names were dropped. Table 1 "
        "summarizes the championship-year files.",
    )
    add_table(
        doc,
        ["Season", "File", "Finishers parsed", "Scoring teams", "Lagged rows", "Notes"],
        rep["audit_rows"],
        "Table 1. Championship-year data audit. All parsed years were contested at Kenilworth Park. Lagged rows enter the estimation sample except 2025, which was held out.",
    )
    body(
        doc,
        "Only varsity five-kilometer championship finishers entered the panel. Models include only schools "
        "with five recorded finishers, so a fifth-runner summary exists by construction. When grade was "
        "labeled, finishers outside grades 9–12 were dropped before standardization. Times outside 14–45 "
        "minutes were dropped, and extras were capped at seven per school. Within each championship, "
        "retained varsity boys were converted to z-scores relative to that year’s field, with higher values "
        "indicating slower relative time. Predictors came from the prior available Kenilworth championship "
        "for the same school. Lag gaps were capped at two years so that 2021 could lag 2019 after the "
        "omitted 2020 meet, and longer skips were excluded. A recovered 2018 varsity dump allows 2019 to "
        "lag 2018. The 2025 championship was reserved as a temporal holdout and was not used for "
        "specification choice.",
    )
    body(
        doc,
        f"The estimation sample comprises {n} team-seasons from {n_schools} schools (2019 and 2021–2024). "
        f"Of {scoring_hold} scoring teams in 2025, {hold_n} retained a usable prior championship "
        f"(Figure 1, Table 2). Prior first- and fifth-runner z-scores were strongly collinear "
        f"(r = {_r(corr)}), so later tests distinguish overlapping summaries rather than "
        f"independent constructs.",
    )
    add_figure(
        doc,
        paths["f1"],
        "Figure 1. Season-by-season counts of parsed varsity finishers, scoring teams (five or more finishers), and lagged analytic rows.",
        width=6.3,
    )
    add_table(
        doc,
        ["Characteristic", "Value"],
        rep["sample_rows"],
        "Table 2. Characteristics of the boys’ estimation sample (fit rows with five or more finishers). Pack gap is fifth-runner z minus first-runner z.",
    )

    subhead(doc, "Statistical Analysis")
    body(
        doc,
        f"Because the same schools appear in several years, models were estimated by ordinary least squares "
        f"with school-clustered standard errors. With {n_schools} clusters, those intervals are treated as "
        f"exploratory descriptors (Cameron & Miller, 2015). Nested models of team place, and of team score "
        f"as a continuous check on the scoring rule, added lagged first-runner z, lagged fifth-runner z, "
        f"lagged pack gap, both z-scores, and conference membership as an exploratory term. Pack gap is a "
        f"linear transform of the two z-scores, so a first-runner-plus-pack-gap specification was not "
        f"estimated separately. Ordered logit models of place were used to check that treating place as "
        f"continuous did not drive the ranking. Akaike information criterion (AIC) differences below about "
        f"2 were not treated as decisive "
        f"(Akaike, 1974).",
    )
    body(
        doc,
        "The confirmatory comparison is 2025 holdout root-mean-square error (RMSE) and mean absolute "
        "error (MAE) for fifth-runner versus first-runner relative time, and versus enrollment and "
        "sector alone. Rank correlation is secondary. Expanding-window rolling-origin error and "
        "leave-one-school-out prediction error provide additional out-of-sample checks. A wild "
        "cluster bootstrap that flips residual signs by school (999 draws) provides a small-sample "
        "interval for the fifth-runner coefficient. A permutation test reassigns lagged fifth-runner z "
        "across team-seasons to test whether the in-sample association is consistent with noise. A "
        "within-team swap test exchanges first- and fifth-runner labels and asks whether those labels "
        "are exchangeable under collinearity; it is not a test of causal superiority. Sensitivities "
        "dropped enrollment, restricted the sample to one-year lags, or replaced z-scores with raw "
        "prior times. An exploratory girls’ panel used the same construction, with 2024 held out "
        "because 2025 had no girls rows with a prior fifth-runner z. "
        f"Analyses were conducted in Python {soft.get('python', '3')} with pandas, NumPy, and "
        f"statsmodels.",
    )

    heading(doc, "Results")
    subhead(doc, "Holdout Evaluation")
    body(
        doc,
        f"On the reserved 2025 championship ({hold_n} of {scoring_hold} scoring teams with a usable prior "
        f"result), prior fifth-runner relative time reduced place error relative to enrollment and sector "
        f"alone and relative to prior first-runner relative time. It was not the lowest-error specification "
        f"among all nested models. The fifth-runner model achieved Spearman ρ = {_r(hd['spearman'])} and "
        f"RMSE = {_r(hd['rmse'])} places (Figure 2), compared with control RMSE of {_r(hc['rmse'])} and "
        f"first-runner RMSE of {_r(hs['rmse'])} (Table 3).",
    )
    body(
        doc,
        f"Pack-gap RMSE was {_r(hg['rmse'])}. Both z-scores reduced error to {_r(hb['rmse'])}, and adding "
        f"conference membership reduced it further to {_r(hw['rmse'])}. "
        + (
            f"A paired bootstrap interval for the fifth-minus-first RMSE difference was {_r(dms['observed'])} "
            f"(95% CI {_r(dms['ci_low'])} to {_r(dms['ci_high'])}), and the interval versus controls was "
            f"{_r(dmc['observed'])} (95% CI {_r(dmc['ci_low'])} to {_r(dmc['ci_high'])}). "
            if dms.get("observed") is not None
            else ""
        )
        + (
            f"Across {roll['n_years']} rolling-origin test seasons, mean place RMSE was "
            f"{_r(roll['mean_rmse_depth'])} for fifth-runner z and {_r(roll['mean_rmse_star'])} for "
            f"first-runner z, with fifth-runner error lower in {roll['depth_beats_star_years']} of "
            f"{roll['n_years']} years. "
            if roll.get("n_years")
            else ""
        )
        + f"Team-score holdout errors ranked similarly. A single later championship cannot separate "
        f"RMSE differences of a few tenths of a place with much precision, and the paired interval for "
        f"fifth-versus-first error included zero.",
    )
    add_table(
        doc,
        ["Model", "Holdout RMSE", "Holdout MAE", "Spearman ρ"],
        rep["hold_table"],
        f"Table 3. Reserved 2025 holdout for team place (n = {hold_n}). RMSE and MAE are in places.",
    )
    add_figure(
        doc,
        paths["f4"],
        "Figure 2. Observed versus predicted 2025 team place from the fifth-runner model. The dashed line marks a perfect numeric match. Lower place is better.",
        width=5.4,
    )

    subhead(doc, "In-Sample Associations")
    vif_s = ""
    if isinstance(vif.get("lag_depth_z"), dict) and "vif" in vif["lag_depth_z"]:
        pcorr_s = _r(pcorr) if pcorr is not None else "n/a"
        vif_s = (
            f" Variance inflation in the joint model was high (variance inflation factor [VIF] ≈ {_r(vif['lag_depth_z']['vif'], 1)} "
            f"for fifth-runner z, and partial correlation after controls was ≈ {pcorr_s}), so the two "
            f"coefficients should not be read as separable pack-depth and star-power effects."
        )
    body(
        doc,
        f"In the estimation sample, the fifth-runner model also showed higher fit for team place "
        f"(Table 4). Enrollment and sector alone explained {_pct(pc['rsquared'])}% of place variation. "
        f"Relative to charter schools, private-school team-seasons were associated with a "
        f"{_r(abs(pc['params']['C(sector)[T.private]']))}-place improvement (lower place is better, "
        f"95% CI {_r(pc['ci_low']['C(sector)[T.private]'])} to "
        f"{_r(pc['ci_high']['C(sector)[T.private]'])}), which is consistent in direction with the "
        f"private-school advantage reported by Johnson et al. (2019) in high-school football. Adding "
        f"first-runner z raised R² to {_r(ps['rsquared'])}, while adding fifth-runner z raised R² to "
        f"{_r(pd_['rsquared'])} and produced the lowest place-model AIC. Each one-standard-deviation slower "
        f"fifth runner was associated with {_r(pd_['params']['lag_depth_z'])} worse places (cluster 95% CI "
        f"{_r(pd_['ci_low']['lag_depth_z'])} to {_r(pd_['ci_high']['lag_depth_z'])}, wild-cluster interval "
        f"{_r(wild['ci_low'])} to {_r(wild['ci_high'])}). Pack gap added substantially less "
        f"(R² = {_r(pg['rsquared'])}, Figure 3), so Falk et al.’s (2022) within-race packing construct "
        f"does not appear here as lagged spread.{vif_s}",
    )
    body(
        doc,
        f"In the joint place model, the fifth-runner coefficient remained larger "
        f"({_r(pb['params']['lag_depth_z'])}) than the first-runner coefficient "
        f"({_r(pb['params']['lag_star_z'])}), but the predictors were highly collinear (r = {_r(corr)}). "
        f"The AIC gap between the fifth-runner-only model and the two-predictor model was "
        f"{_r(m['delta_aic']['both'], 1)}, which is not decisive under Akaike’s (1974) criterion. Ordered "
        f"logit ranked those two specifications together and ahead of first-runner z. Figure 4 shows both "
        f"z-scores against next championship place. Leave-one-school-out prediction RMSE for the "
        f"fifth-runner place model was {_r(loo_rmse)} places.",
    )
    add_table(
        doc,
        ["Model", "N", "R²", "Adj. R²", "AIC", "ΔAIC", "Focal coef (SE)", "95% CI"],
        rep["place_table"],
        f"Table 4. Ordinary least squares models of next championship boys’ team place with school-clustered standard errors ({n_schools} schools). ΔAIC is relative to the fifth-runner model.",
    )
    add_figure(
        doc,
        paths["f2"],
        "Figure 3. Prior pack gap versus next championship team place in the estimation sample, with an ordinary least squares fit. Lower place is better.",
        width=5.4,
    )
    add_figure(
        doc,
        paths["f3"],
        "Figure 4. Prior first-runner z (left) and fifth-runner z (right) versus next championship team place. Lower z-scores indicate faster relative time against the year’s field.",
        width=6.4,
    )
    body(
        doc,
        f"Team score, the championship total of five places, produced the same qualitative ranking. "
        f"Fifth-runner relative time raised in-sample fit more than first-runner relative time "
        f"(R² = {_r(sd['rsquared'])} versus {_r(ss['rsquared'])}), while controls explained "
        f"{_pct(sc['rsquared'])}%. In the joint score model, the fifth-runner interval excluded zero and "
        f"the first-runner interval did not. Leave-one-school-out fifth-runner coefficients stayed in a "
        f"narrow band ({_r(loo_c['min'])} to {_r(loo_c['max'])}). A permutation test of the fifth-runner "
        f"place model was inconsistent with noise alone (p = {_r(perm['p'], 3)}), but the "
        f"within-team swap test did not separate the two labels (p = {_r(swap['p'], 2)}). Conference "
        f"membership was not distinguishable from zero in the estimation sample, even though it entered "
        f"the lowest-error 2025 forecast. An exploratory girls’ panel ({gc['n']} team-seasons, 2024 holdout "
        f"n = {gh['controls']['n']}) ranked first-runner z higher in sample (R² = {_r(gs['rsquared'])} "
        f"versus {_r(gd['rsquared'])}), while fifth-runner z retained slightly lower holdout error. That "
        f"sample is too small to establish effect modification by sex.",
    )

    heading(doc, "Discussion")
    body(
        doc,
        "The confirmatory comparison is whether prior fifth-runner relative time predicts next "
        "championship boys’ team place more accurately than prior first-runner relative time after "
        "enrollment and sector adjustment. In this Kenilworth panel it does, both in the estimation "
        "years and on the reserved 2025 championship. That ranking fits the five-scorer rule, because "
        "the fifth finisher is the last place that enters the team total, but it is not evidence that "
        "pack depth has been isolated from overall team speed. First- and fifth-runner z-scores were "
        "highly collinear, variance inflation in the joint model was high, and the within-team swap "
        "test did not distinguish the two labels once both times were known. School size and sector "
        "already accounted for a large share of place variation, which matches the private-school "
        "advantage Johnson, Manwell, and Scott (2019) reported in another interscholastic setting. "
        "Adding the fifth runner still improved the fit and reduced holdout error relative to those "
        "controls.",
    )
    body(
        doc,
        "The same collinearity also limits how far the packing literature can be applied here. Falk, "
        "Larson, and DeBeliso (2022) measured packing as within-race clustering among NCAA teammates, "
        "and Hanley (2015, 2016) and Finley and Fountain (2023) likewise analyze behavior after the "
        "gun: pacing, packing, and movement through a championship field. If those within-race "
        "constructs were interchangeable with two lagged finish times, the spread between first and "
        "fifth times should have been the more informative predictor of next place. Pack gap was not, "
        "while the level of the fifth runner was. A lagged finish time could mean a thin roster, a poor "
        "championship day, or a deeper opposing field, and without checkpoint splits those mechanisms "
        "cannot be separated. Falk et al. (2022) imply that last year’s packing, proxied by spread, "
        "should travel, and the data do not support that implication. Fast programs in this "
        "association tend to be fast at both ends of the scoring five, so collinearity of about 0.85 "
        "is an empirical fact rather than a software warning.",
    )
    body(
        doc,
        "The sports-forecasting literature supplies the evaluation design rather than a competing "
        "model. Stefani (1980), Kvam and Sokol (2006), and Pasteur and Janning (2011) rate teams from "
        "repeated contests against identified opponents. A District of Columbia championship supplies "
        "one race per year, so holding out 2025 follows from that literature’s insistence on "
        "out-of-sample judgment rather than from an attempt to build a comparable rating system on two "
        "lagged times. Nested specifications that used both z-scores, or conference membership, "
        "reduced 2025 error further. Extra predictors can improve a one-meet forecast without changing "
        "the confirmatory question. The paired bootstrap interval for fifth-versus-first RMSE included "
        "zero, so the holdout should be read as leaning the same way as the estimation sample, not as "
        "a precise ranking of nested models.",
    )
    body(
        doc,
        "The exploratory girls’ panel should not be treated as a test of Hanley’s (2016) sex "
        "differences in elite packing. That study concerns world-level marathon fields, while the "
        "girls’ sample here is a small set of schools and seasons, with 2024 held out because 2025 did "
        "not yield lagged fifth-runner rows. First-runner z ranked higher in sample, and fifth-runner "
        "z retained slightly lower holdout error. Either pattern is consistent with sampling variation. "
        "The panel does not support a claim that fifth-runner form is generally better, or that the "
        "boys’ ranking is sex-specific.",
    )
    body(
        doc,
        f"Most of the limits come from the sample. Parseable listings come from one district and one "
        f"course ({n} lagged team-seasons, {n_schools} schools). Schools that do not field five "
        "finishers in consecutive available championships never enter, so stable programs are over-"
        "represented. Some lags skip a year. Z-scores shift when the opposing field shifts. Private "
        "and charter enrollment is mostly a constant. Place is an ordered rank treated as continuous, "
        "though ordered logit and team score produced the same qualitative ranking. Conference "
        "membership is a binary schedule proxy. Only the holdout fifth-versus-first error comparison "
        "was locked before inspection of the 2025 results. The design is observational and does not "
        "identify causal effects of roster construction.",
    )
    body(
        doc,
        "Within those bounds, the implication is local. For District of Columbia boys’ championships "
        "at Kenilworth, last year’s fifth-runner relative time is the better of the two public "
        "individual summaries for next team place once school size and sector are in the model. The "
        "same results page does not show that packing caused that place, and the ranking need not "
        "travel to a classified state meet on a different course. Repeating the same first-versus-fifth "
        "comparison in an association that uses classifications would test whether the "
        "fifth-scorer advantage is a feature of the scoring rule or of this particular open field.",
    )

    heading(doc, "Conclusion")
    body(
        doc,
        "Among schools with an eligible prior Kenilworth championship, prior fifth-runner relative "
        "time predicted subsequent District of Columbia boys’ team place and team score more closely "
        "than prior first-runner relative time after adjustment for enrollment and sector, and it "
        "produced smaller errors on a reserved later championship than either school characteristics "
        "or the first runner. Specifications that used both runners, or conference membership, were "
        "still more accurate on that later meet, but the two runner summaries were too collinear to "
        "support a separable pack-depth interpretation. The ranking applies to this association and "
        "course, and it does not establish that the fifth scorer is always the statistic that matters.",
    )

    heading(doc, "Data Availability")
    body(
        doc,
        f"Replication code, cleaned team-season files, the season audit, and the {n}-row analytic "
        f"sample are available at {REPO}. The underlying listings are public championship results from "
        f"the District of Columbia State Athletic Association, MileSplit, and M&D Timing.",
    )

    heading(doc, "References")
    refs = [
        "Akaike, H. (1974). A new look at the statistical model identification. IEEE Transactions on Automatic Control, 19(6), 716–723. https://doi.org/10.1109/TAC.1974.1100705",
        "Cameron, A. C., & Miller, D. L. (2015). A practitioner’s guide to cluster-robust inference. Journal of Human Resources, 50(2), 317–372. https://doi.org/10.3368/jhr.50.2.317",
        "District of Columbia Public Schools. (2026). DCPS data set – enrollment (audit files). https://dcps.dc.gov/node/1018342",
        "District of Columbia State Athletic Association. (2025). 2025 cross country bulletin. https://www.dcsaasports.com/boys-cross-country/",
        "Falk, J., Larson, A., & DeBeliso, M. (2022). Pack running among female NCAA collegiate cross country teams. European Journal of Physical Education and Sport Science, 9(1). https://doi.org/10.46827/ejpe.v9i1.4484",
        "Finley, P. S., & Fountain, J. J. (2023). An examination of race strategies in NCAA cross country championship events. Journal of Coaching and Sports Science, 2(1), 1–11. https://doi.org/10.58524/jcss.v2i1.210",
        "Hanley, B. (2015). Pacing profiles and pack running at the IAAF World Half Marathon Championships. Journal of Sports Sciences, 33(11), 1189–1195. https://doi.org/10.1080/02640414.2014.988743",
        "Hanley, B. (2016). Pacing, packing and sex-based differences in Olympic and IAAF World Championship marathons. Journal of Sports Sciences, 34(18), 1675–1681. https://doi.org/10.1080/02640414.2015.1132841",
        "Johnson, J. E., Manwell, A. K., & Scott, B. F. (2019). An examination of competitive balance within interscholastic football. Journal of Amateur Sport, 5(1), 21–49. https://doi.org/10.17161/jas.v5i1.6708",
        "Kvam, P., & Sokol, J. S. (2006). A logistic regression/Markov chain model for NCAA basketball. Naval Research Logistics, 53(8), 788–803. https://doi.org/10.1002/nav.20165",
        "M&D Timing. (2025). DCSAA Cross Country Championships 2025 results. https://results.mdtimingllc.com/meets/58668",
        "MileSplit. (2026). DCSAA Cross Country Championships results index. https://dc.milesplit.com/",
        "National Federation of State High School Associations. (2025). 2025–26 NFHS track and field and cross country rules book. NFHS.",
        "Pasteur, R. D., & Janning, M. C. (2011). Monte Carlo simulation for high school football playoff seed projection. Journal of Quantitative Analysis in Sports, 7(2), Article 11. https://doi.org/10.2202/1559-0410.1330",
        "Stefani, R. T. (1980). Improved least squares football, basketball, and soccer ratings. IEEE Transactions on Systems, Man, and Cybernetics, 10(2), 116–123. https://doi.org/10.1109/TSMC.1980.4308442",
    ]
    for ref in refs:
        hanging(doc, ref)

    doc.save(DOCX)
    return DOCX


def main():
    jpgs = make_figures()
    pngs = export_png(jpgs)
    out = build_docx(jpgs)
    print("wrote", out)
    pdf = export_pdf(out, PDF)
    print("wrote", pdf)
    for k, v in pngs.items():
        print(k, v)


if __name__ == "__main__":
    main()
